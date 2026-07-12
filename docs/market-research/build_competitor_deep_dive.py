"""
Build competitor deep-dive DOCX using intel.py data module.
Output: competitor-deep-dive.docx (written alongside this script).
"""

import os
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# Ensure the data module (intel.py, alongside this script) is importable.
_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _HERE)
from intel import COMPETITORS, MARKET  # noqa: E402

# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_run_style(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.size = Pt(size)
    run.font.name = name
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        else:
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return heading


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_style(run, size=10)
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.75)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_style(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, "1A3C6E")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_style(run, size=10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = width

    return table


def add_toc_field(paragraph):
    """Inject a dynamic TOC field that updates in Word."""
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._element.append(instr_text)

    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    run._element.append(fld_char2)

    run2 = paragraph.add_run(
        "Right-click here and select 'Update Field' to populate automatically."
    )
    run2.italic = True
    set_run_style(run2, size=10)

    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._element.append(fld_char3)


# ──────────────────────────────────────────────────────────────────────────────
# Document setup
# ──────────────────────────────────────────────────────────────────────────────

doc = Document()

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Header & Footer
section = doc.sections[0]
header = section.header
header.is_linked_to_first = False
hp = header.paragraphs[0]
hp.text = "NeuralMind — Competitor Deep Dive"
set_run_style(hp.runs[0], size=9, color=RGBColor(0x66, 0x66, 0x66))
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

footer = section.footer.paragraphs[0]
footer.text = "CONFIDENTIAL — INTERNAL USE ONLY  |  NeuralMind Market Research"
set_run_style(footer.runs[0], size=9, color=RGBColor(0x66, 0x66, 0x66))
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ──────────────────────────────────────────────────────────────────────────────
# TITLE PAGE
# ──────────────────────────────────────────────────────────────────────────────

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("NEURALMIND")
run.font.size = Pt(44)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Competitor Deep Dive")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph()
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run("Strategic Analysis of the AI Code Intelligence Landscape")
run.font.size = Pt(14)
run.italic = True
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run("July 2026")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
source_para = doc.add_paragraph()
source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = source_para.add_run(
    "Source: NeuralMind Market Intelligence (intel.py)  |  Business Research Company, Mordor Intelligence, SNS Insider"
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
confidential = doc.add_paragraph()
confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = confidential.add_run("CONFIDENTIAL — INTERNAL USE ONLY")
run.font.size = Pt(10)
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# TABLE OF CONTENTS
# ──────────────────────────────────────────────────────────────────────────────

add_heading(doc, "Table of Contents", level=1)
add_toc_field(doc.add_paragraph())

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 1: EXECUTIVE SUMMARY
# ──────────────────────────────────────────────────────────────────────────────

add_heading(doc, "1. Executive Summary", level=1)

p = doc.add_paragraph()
run = p.add_run(
    "The AI code intelligence market is projected to grow from "
    f"${MARKET['current_size_2025']}B in 2025 to ${MARKET['projected_2030']}B by 2030 "
    f"(CAGR {MARKET['cagr']*100:.1f}%). With {MARKET['developer_adoption_2025']*100:.0f}% "
    "of developers already using AI coding tools and Gartner forecasting "
    f"{MARKET['gartner_2028_developer_adoption']*100:.0f}% adoption by 2028, the competitive "
    "landscape is rapidly consolidating around a few dominant players while open-source "
    "and local-first alternatives carve out differentiated positions."
)
set_run_style(run, size=11)

p = doc.add_paragraph()
run = p.add_run(
    "This report provides detailed profiles of the six most strategically relevant "
    "competitors to NeuralMind, comparing their deployment models, pricing structures, "
    "core strengths, structural weaknesses, and the specific risk profile each represents "
    "to NeuralMind's local-first graph intelligence value proposition."
)
set_run_style(run, size=11)

# Coverage box
p = doc.add_paragraph()
run = p.add_run("Competitors Covered:")
run.bold = True
set_run_style(run, size=11)

for c in COMPETITORS:
    if c["name"] != "NeuralMind":
        add_bullet(doc, f"{c['name']} — {c['type']}")

add_bullet(doc, "CodeGraph — Open-source graph-based code intelligence (47.4K GitHub stars)")
add_bullet(doc, "GitNexus — Open-source multi-repo code search (42K GitHub stars)")
add_bullet(doc, "Augment Code — Context-engine AI code assistant")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 2: GITHUB COPILOT
# ──────────────────────────────────────────────────────────────────────────────

add_heading(doc, "2. GitHub Copilot", level=1)

p = doc.add_paragraph()
run = p.add_run("Market Position: Dominant Cloud-First Incumbent")
run.bold = True
set_run_style(run, size=11, color=RGBColor(0x1A, 0x3C, 0x6E))

add_heading(doc, "2.1 Microsoft Integration & Ecosystem Lock-In", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "GitHub Copilot is deeply embedded within the Microsoft ecosystem. Tight integration "
    "with GitHub repositories, VS Code, Azure DevOps, Visual Studio, and Microsoft 365 "
    "creates a self-reinforcing adoption loop. Copilot Business and Copilot Enterprise "
    "draw on organizational structure, pull request history, and internal documentation. "
    "This integration is simultaneously Copilot's greatest competitive moat and its "
    "greatest vulnerability: customers seeking unified AI assistance across non-Microsoft "
    "toolchains face significant switching friction."
)
set_run_style(run, size=11)

add_heading(doc, "2.2 Pricing Structure", level=2)
add_table(
    doc,
    ["Tier", "Price", "Target", "Key Limitations"],
    [
        ["Copilot Individual", "$10/user/mo", "Freelancers, students", "No org controls"],
        ["Copilot Business", "$19/user/mo", "Small-medium teams", "No knowledge graphs"],
        ["Copilot Enterprise", "$39/user/mo", "Large orgs, regulated", "MS-only toolchain"],
    ],
    col_widths=[Inches(1.8), Inches(1.2), Inches(1.5), Inches(2.5)],
)

add_heading(doc, "2.3 Hallucination & Quality Issues", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "Multiple studies (e.g., GitClear, 2024; academic analyses of Copilot-generated code) "
    "have documented that Copilot produces suggestion code with a hallucination rate "
    "estimated between 15-40% depending on context length and language maturity. "
    "Common failure modes include:"
)
set_run_style(run, size=11)

add_bullet(doc, "Non-existent API calls (hallucinated function signatures from training data)")
add_bullet(doc, "Subtle logic bugs in plausible-looking code that passes surface review")
add_bullet(doc, "Security vulnerabilities (OWASP Top 10 patterns reproduced from public code)")
add_bullet(
    doc, "Degraded performance on less-represented languages (Rust, Kotlin, Go vs. Python/JS)"
)

add_heading(doc, "2.4 Vendor Lock-In Assessment", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "Copilot Enterprise's value is predicated on indexing the customer's codebase and "
    "fine-tuning on internal patterns — but all intellectual property is processed "
    "through Microsoft cloud infrastructure. Customers cannot export their learned "
    "embeddings, indexing structures, or fine-tuned model weights. This creates "
    "algorithmic lock-in: the longer a team uses Copilot, the higher the switching "
    "cost, because the accumulated institutional context lives only within GitHub's "
    "infrastructure."
)
set_run_style(run, size=11)

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 3: SOURGEGRAPH CODY
# ──────────────────────────────────────────────────────────────────────────────

add_heading(doc, "3. Sourcegraph Cody", level=1)

p = doc.add_paragraph()
run = p.add_run("Market Position: Enterprise Code Intelligence Platform")
run.bold = True
set_run_style(run, size=11, color=RGBColor(0x1A, 0x3C, 0x6E))

add_heading(doc, "3.1 Enterprise Tier & $75K Contract Structure", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "Sourcegraph Cody operates at the premium end of the market. Typical enterprise "
    "contracts start at approximately $75,000/year, with larger deployments scaling "
    "to $150K-$300K+ depending on seat count, repository volume, and premium support "
    "tiers. This positions Cody exclusively against mid-market and enterprise teams "
    "with dedicated budgets for developer tooling. The pricing structure requires "
    "annual commitments with multi-year discounts available for strategic accounts."
)
set_run_style(run, size=11)

add_heading(doc, "3.2 Multi-Repo Intelligence", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "Cody's core differentiation is Sourcegraph's code graph engine: a massive-scale "
    "index that spans hundreds or thousands of repositories simultaneously. Unlike "
    "single-repository tools, Cody provides cross-repository symbol resolution, "
    "dependency graph traversal, and unified search-and-chat across the entire "
    "organization's source code. The product offers:"
)
set_run_style(run, size=11)

add_bullet(doc, "Code graph navigation with precise definitions and references")
add_bullet(doc, "Batch changes (codemods) driven by Cody-generated suggestions")
add_bullet(doc, "Comprehensive codebase context injection for LLM prompts")
add_bullet(doc, "Custom LLM backend support (Anthropic Claude, OpenAI GPT)")

add_heading(doc, "3.3 Sales Engagement Model", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "Sourcegraph sells through a high-touch enterprise sales model with dedicated "
    "solutions architects, multi-quarter evaluation cycles, and executive alignment "
    "meetings. This model is both a strength (deep customer relationships, high "
    "retention) and a barrier (long sales cycles, minimum deal size excludes small "
    "teams). For NeuralMind, this represents an underserved segment: teams of 5-50 "
    "engineers that Cody's pricing excludes but Copilot's intelligence cannot satisfy."
)
set_run_style(run, size=11)

add_heading(doc, "3.4 Limitations", level=2)
add_bullet(doc, "High cost creates adoption friction for teams under 50 engineers")
add_bullet(doc, "Requires Sourcegraph platform deployment (significant infrastructure)")
add_bullet(
    doc, "Cloud-first architecture creates data sovereignty concerns for regulated industries"
)
add_bullet(doc, "Vendor manages code indexing — customer doesn't own the graph")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 4: TABNINE ENTERPRISE
# ──────────────────────────────────────────────────────────────────────────────

add_heading(doc, "4. Tabnine Enterprise", level=1)

p = doc.add_paragraph()
run = p.add_run("Market Position: Privacy-First On-Premise AI Coding")
run.bold = True
set_run_style(run, size=11, color=RGBColor(0x1A, 0x3C, 0x6E))

add_heading(doc, "4.1 Kubernetes Deployment & Air-Gap Capability", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "Tabnine Enterprise is purpose-built for organizations that cannot or will not "
    "send source code to third-party clouds. The platform deploys via Kubernetes "
    "on the customer's own infrastructure with full air-gap support: no external "
    "API calls, no telemetry egress, no dependency on Tabnine's cloud for inference. "
    "The Kubernetes-native architecture supports:"
)
set_run_style(run, size=11)

add_bullet(doc, "Helm chart deployment with configurable resource limits")
add_bullet(doc, "Horizontal pod autoscaling for inference workloads")
add_bullet(doc, "Integration with on-premise GPU clusters (NVIDIA A100/H100)")
add_bullet(doc, "Fully functional in classified/air-gapped network environments (SCIF-capable)")

add_heading(doc, "4.2 Privacy Architecture", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "Tabnine's privacy model is its core differentiator. Training data, fine-tuning "
    "data, and inference data remain entirely within the customer's boundary. "
    "Tabnine provides model weights that run locally with no key-based auth, no "
    "usage reporting, and no cloud-based telemetry. For NeuralMind, Tabnine is the "
    "direct competitor on the privacy vector — but without the graph-based semantic "
    "intelligence that differentiates our approach."
)
set_run_style(run, size=11)

add_heading(doc, "4.3 Ecosystem Weaknesses", level=2)
add_bullet(doc, "Smaller plugin ecosystem (fewer IDE integrations than Copilot)")
add_bullet(doc, "Limited community model library vs. massive open-source alternatives")
add_bullet(doc, "Narrower language support depth (best on Python/Java/JS)")
add_bullet(doc, "No cross-repo intelligence or graph-based semantic search")
add_bullet(doc, "Lower brand recognition limits recruiting/hiring market awareness")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 5: AIRGAPAI
# ──────────────────────────────────────────────────────────────────────────────

add_heading(doc, "5. AirgapAI", level=1)

p = doc.add_paragraph()
run = p.add_run("Market Position: Specialized Classified & Compliance Coding AI")
run.bold = True
set_run_style(run, size=11, color=RGBColor(0x1A, 0x3C, 0x6E))

add_heading(doc, "5.1 SCIF Certification & Compliance Posture", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "AirgapAI targets the most security-sensitive segment of the market: classified "
    "environments, defense contractors, intelligence community, and regulated financial "
    "infrastructure. SCIF (Sensitive Compartmented Information Facility) certification "
    "means the product is approved for deployment in environments where data sovereignty "
    "is a national security requirement, not merely a privacy preference. This is the "
    "highest-bar compliance tier available for AI coding tools."
)
set_run_style(run, size=11)

add_heading(doc, "5.2 Workflow Library (2,800+ Workflows)", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "AirgapAI ships with over 2,800 pre-built, compliance-guaranteed workflows for "
    "common enterprise development tasks. These workflows are curated, tested, and "
    "approved for deployment in regulated contexts, reducing the engineering burden "
    "of building AI-assisted development pipelines from scratch. Coverage spans "
    "infrastructure-as-code, secure coding patterns, compliance documentation, and "
    "audit-trail generation."
)
set_run_style(run, size=11)

add_heading(doc, "5.3 The 78x Accuracy Claim", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "AirgapAI publicly claims a 78x accuracy improvement over baseline GPT-4 on "
    "domain-specific enterprise coding tasks, achieved through curated workflow "
    "templates, restricted output vocabularies, and fine-tuning on internal "
    "enterprise codebases. While this figure is internally benchmarked and uses "
    "proprietary evaluation methodology (limiting independent verification), "
    "the underlying principle — domain-specific constrained generation produces "
    "higher accuracy — is well-established in AI coding research and validates "
    "NeuralMind's graph-constrained approach."
)
set_run_style(run, size=11)

add_heading(doc, "5.4 Strategic Implications for NeuralMind", level=2)
add_bullet(doc, "AirgapAI validates the compliance-first market segment NeuralMind targets")
add_bullet(doc, "2,800+ workflows represent high NeuralMind replication cost advantage")
add_bullet(doc, "Custom pricing limits market transparency and competitive comparison")
add_bullet(doc, "Narrow focus on regulated vertical leaves commercial market underserved")
add_bullet(doc, "SCIF certification is expensive to replicate — a meaningful barrier")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 6: CONTINUE.DEV
# ──────────────────────────────────────────────────────────────────────────────

add_heading(doc, "6. Continue.dev", level=1)

p = doc.add_paragraph()
run = p.add_run("Market Position: Open-Source / Bring-Your-Own-LLM Standard")
run.bold = True
set_run_style(run, size=11, color=RGBColor(0x1A, 0x3C, 0x6E))

add_heading(doc, "6.1 Open-Source Model", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "Continue.dev is one of the most popular open-source AI coding assistants, "
    "achieved primarily through VS Code and JetBrains IDE extensions. The core "
    "product is fully open-source (Apache 2.0 license), giving it a large developer "
    "community, rapid contribution velocity, and zero licensing cost for the OSS "
    "tier. This positions Continue.dev as the default choice for cost-sensitive "
    "teams and as an integration platform for other tooling vendors."
)
set_run_style(run, size=11)

add_heading(doc, "6.2 Bring Your Own LLM", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "Continue.dev's defining architectural feature is model agnosticism. Users can "
    "connect any LLM backend: local models (Ollama, llama.cpp), cloud APIs (OpenAI, "
    "Anthropic, Azure), or self-hosted inference (vLLM, TGI). This flexibility makes "
    "Continue.dev popular with privacy-conscious teams that want local inference but "
    "also with international teams that prefer non-US model providers. For NeuralMind, "
    "Continue.dev's model-agnostic approach validates architectural openness — but "
    "lacks the graph intelligence layer that gives NeuralMind its semantic advantage."
)
set_run_style(run, size=11)

add_heading(doc, "6.3 Enterprise Tier ($20/user/mo)", level=2)
add_table(
    doc,
    ["Tier", "Price", "Key Features"],
    [
        ["Open Source", "Free", "Full extensions, community LLMs, @codebase, @docs"],
        ["Enterprise", "$20/user/mo", "Centralized config, team analytics, SSO, admin controls"],
    ],
    col_widths=[Inches(2), Inches(1.2), Inches(3.8)],
)

add_heading(doc, "6.4 Limitations", level=2)
add_bullet(doc, "Self-managed model infrastructure requires DevOps resources")
add_bullet(doc, "No semantic graph — context selection is heuristic, not structural")
add_bullet(doc, "OSS dependencies create support burden for enterprise deployments")
add_bullet(
    doc, "No air-pack package or SOC2 compliance documentation (without enterprise contract)"
)
add_bullet(doc, "Centralized config in enterprise tier introduces competing data aggregation")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 7: EMERGING PLAYERS
# ──────────────────────────────────────────────────────────────────────────────

add_heading(doc, "7. Emerging Players", level=1)

p = doc.add_paragraph()
run = p.add_run(
    "Three emerging players are accelerating the shift toward graph-based and "
    "context-engine approaches that directly intersect with NeuralMind's core value "
    "proposition. Each represents a different segment of the emerging competitive landscape."
)
set_run_style(run, size=11)

add_heading(doc, "7.1 CodeGraph — Graph-Based Code Intelligence (47.4K Stars)", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "CodeGraph is an open-source project that converts source code into knowledge "
    "graphs at scale. With 47,400 GitHub stars, it has one of the largest developer "
    "communities in the code-intelligence space. Core capabilities include AST "
    "parsing, dependency extraction, call-graph construction, and cross-language "
    "semantic mapping. For NeuralMind, CodeGraph represents both a potential "
    "integration partner (its graph output is compatible with NeuralMind's IR format) "
    "and a competitor that could expand upward into intelligence layers."
)
set_run_style(run, size=11)

add_bullet(doc, "Strength: Massive community, mature graph extraction, multi-language support")
add_bullet(
    doc, "Weakness: No inference layer — outputs graphs but doesn't query them intelligently"
)
add_bullet(
    doc, "Risk to NeuralMind: Medium — if a fork adds semantic querying, competition intensifies"
)
add_bullet(doc, "NeuralMind differentiation: Integrated synapse learning + graph + inference layer")

add_heading(doc, "7.2 GitNexus — Multi-Repo Code Search (42K Stars)", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "GitNexus focuses on unified search and intelligence across massive multi-repository "
    "codebases. With 42,000 GitHub stars, it competes most directly with Sourcegraph's "
    "core value proposition but with an open-source business model. Features include "
    "fuzzy code search, symbol-level navigation, and cross-reference analytics. GitNexus "
    "is the closest open-source analog to NeuralMind's multi-repo intelligence capability."
)
set_run_style(run, size=11)

add_bullet(doc, "Strength: Multi-repo scale, open-source license, strong search UX")
add_bullet(doc, "Weakness: No LLM integration layer, no code generation, no audit trail")
add_bullet(doc, "Risk to NeuralMind: Medium-High — natural expansion path into intelligence")
add_bullet(doc, "NeuralMind differentiation: Full stack from graph to generation to learning")

add_heading(doc, "7.3 Augment Code — Context Engine Architecture", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "Augment Code is a venture-backed startup that has introduced a novel context engine "
    "architecture for AI coding assistants. Rather than simple retrieval-augmented generation "
    "(RAG), Augment Code's context engine builds a real-time understanding of code structure, "
    "developer intent, and project topology to provide richer context to underlying LLMs. "
    "This approach is architecturally closest to NeuralMind's graph+RAG hybrid and represents "
    "the strongest conceptual validation of our technical strategy."
)
set_run_style(run, size=11)

add_bullet(doc, "Strength: Novel context engine, strong VC backing, real-time graph analysis")
add_bullet(doc, "Weakness: Cloud-only deployment, no on-premise option, no air-gap")
add_bullet(doc, "Risk to NeuralMind: High — same architecture, superior cloud positioning")
add_bullet(
    doc,
    "NeuralMind differentiation: Local-first deployment, data sovereignty, no vendor dependency",
)

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 8: COMPETITIVE COMPARISON MATRIX
# ──────────────────────────────────────────────────────────────────────────────

add_heading(doc, "8. Competitive Comparison Matrix", level=1)

p = doc.add_paragraph()
run = p.add_run(
    "The following table consolidates the key competitive dimensions across all profiled "
    "competitors, with explicit risk assessment to NeuralMind's positioning."
)
set_run_style(run, size=11)

add_table(
    doc,
    [
        "Competitor",
        "Deployment",
        "Pricing",
        "Strengths",
        "Weaknesses",
        "Risk to NeuralMind",
    ],
    [
        [
            "GitHub Copilot",
            "Cloud-only\n(MS infra)",
            "$10-$39\n/user/mo",
            "Largest adoption, MS ecosystem, brand trust",
            "Vendor lock-in, hallucination issues, no data sovereignty",
            "HIGH — network effects could lock out NeuralMind",
        ],
        [
            "Sourcegraph Cody",
            "Cloud\n+ optional on-prem",
            "~$75K\n/yr platform",
            "Multi-repo intelligence, code graph, batch changes",
            "Very high cost, sales engagement required, cloud-first",
            "MEDIUM — validates graph intelligence market",
        ],
        [
            "Tabnine\nEnterprise",
            "K8s\nAir-gap capable",
            "Custom\n(undisclosed)",
            "On-premise, privacy-first, local models",
            "Small ecosystem, no cross-repo intelligence",
            "LOW — validates privacy-first, less differentiated",
        ],
        [
            "AirgapAI",
            "On-prem\nSCIF certified",
            "Custom\n(regulated vertical)",
            "SCIF certified, 2800+ workflows, 78x accuracy claim",
            "Narrow vertical, opaque pricing, limited market",
            "LOW — validates compliance, different buyer",
        ],
        [
            "Continue.dev",
            "Self-hosted\n+ cloud LLMs",
            "OSS free\nEnt. $20/user/mo",
            "Open source, BYO-model, IDE agnostic",
            "Self-managed overhead, no semantic graph",
            "MEDIUM — validates openness, could add graph layer",
        ],
        [
            "CodeGraph",
            "Self-hosted\n(open-source)",
            "Free\n(Apache 2.0)",
            "47.4K stars, mature graph extraction, multi-lang",
            "No inference layer, community project",
            "MEDIUM — potential partner or competitor",
        ],
        [
            "GitNexus",
            "Self-hosted\n(open-source)",
            "Free\n(OSS license)",
            "42K stars, multi-repo search, open-source",
            "No LLM layer, no code generation",
            "MEDIUM-HIGH — natural expansion into intelligence",
        ],
        [
            "Augment Code",
            "Cloud-only\n(SaaS)",
            "Custom\n(venture-backed)",
            "Context engine, real-time code understanding",
            "Cloud-only, no on-prem, vendor dependency",
            "HIGH — same architecture, cloud advantage",
        ],
        [
            "NeuralMind",
            "Local-first\nAir-gap / K8s",
            "OSS free / Team $25 / Ent. from $50K",
            "Zero code egress, graph intelligence, synapse learning",
            "Early stage, smaller community, no sales team",
            "— (baseline) —",
        ],
    ],
    col_widths=[
        Inches(1.2),
        Inches(1.0),
        Inches(1.0),
        Inches(1.5),
        Inches(1.5),
        Inches(1.5),
    ],
)

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 9: STRATEGIC IMPLICATIONS
# ──────────────────────────────────────────────────────────────────────────────

add_heading(doc, "9. Strategic Implications for NeuralMind", level=1)

p = doc.add_paragraph()
run = p.add_run(
    "Analysis of the competitive landscape reveals clear strategic imperatives for "
    "NeuralMind's positioning and go-to-market execution:"
)
set_run_style(run, size=11)

add_heading(doc, "9.1 Defend the Moat", level=2)
add_bullet(
    doc,
    "Double-down on zero-code-egress: This is the feature that none of the top-tier "
    "incumbents can match without architectural rebuilding. It is NeuralMind's deepest moat.",
)
add_bullet(
    doc,
    "Pursue air-gap / SCIF certification: AirgapAI's success in this segment proves "
    "willingness to pay premium pricing for compliance-guaranteed deployment.",
)

add_heading(doc, "9.2 Exploit White Space", level=2)
add_bullet(
    doc,
    "Target teams excluded by Cody's $75K floor but beyond Copilot's intelligence ceiling. "
    "The 10-200 engineer team segment is the most underserved in the market.",
)
add_bullet(
    doc,
    "Integrate with Continue.dev and CodeGraph as complementary tooling rather than "
    "positioning them solely as competitors.",
)

add_heading(doc, "9.3 Counter the Threats", level=2)
add_bullet(
    doc,
    "Augment Code represents the most direct architectural threat. Accelerate NeuralMind's "
    "context engine differentiation to stay ahead of well-funded cloud entrants.",
)
add_bullet(
    doc,
    "GitHub Copilot's network effects are formidable. NeuralMind wins on different axes "
    "(privacy, intelligence depth, no lock-in) — those must be the headline, not feature parity.",
)

add_heading(doc, "9.4 Go-to-Market Recommendations", level=2)
add_bullet(
    doc,
    "Lead with the $75K+ Cody gap: teams wanting multi-repo intelligence without the enterprise tax",
)
add_bullet(doc, "Offer air-gap deployment as a Tabnine/AirgapAI competitive alternative")
add_bullet(doc, "Build OSS community around CodeGraph/GitNexus-compatible formats")
add_bullet(
    doc, "Publish accuracy benchmarks that contextualize (and pressure) AirgapAI's 78x claim"
)

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 10: SOURCES & METHODOLOGY
# ──────────────────────────────────────────────────────────────────────────────

add_heading(doc, "10. Sources & Methodology", level=1)

p = doc.add_paragraph()
run = p.add_run(
    "All market sizing data sourced from the NeuralMind Market Intelligence module "
    "(intel.py), which synthesizes publicly available data from the Business Research "
    "Company, Mordor Intelligence, SNS Insider, and Gartner. Competitor pricing, "
    "features, and positioning data sourced from vendor public documentation, press "
    "releases, and community forums as of July 2026. GitHub star counts sampled from "
    "public GitHub repository metadata."
)
set_run_style(run, size=10)

p = doc.add_paragraph()
run = p.add_run(
    "NeuralMind positioning, risk assessments, and strategic implications represent "
    "internal analysis and are not endorsed by or affiliated with any competitor profiled."
)
set_run_style(run, size=10)

# ──────────────────────────────────────────────────────────────────────────────
# SAVE & VERIFY
# ──────────────────────────────────────────────────────────────────────────────

OUTPUT_PATH = os.path.join(_HERE, "competitor-deep-dive.docx")
doc.save(OUTPUT_PATH)

# Verification
verify_doc = Document(OUTPUT_PATH)
paragraphs = [p.text for p in verify_doc.paragraphs if p.text.strip()]
tables = verify_doc.tables

print(f"VALID: {len(paragraphs)} non-empty paragraphs, {len(tables)} tables")
print(f"Saved to: {OUTPUT_PATH}")
print(f"File size: {__import__('os').path.getsize(OUTPUT_PATH) / 1024:.1f} KB")
