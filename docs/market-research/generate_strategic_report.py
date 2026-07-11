#!/usr/bin/env python3
"""Generate NeuralMind Strategic Report DOCX."""

import sys
sys.path.insert(0, "/home/dtfrost/neuralmind/docs/market-research")

from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from intel import (
    MARKET, COMPETITORS, PILLARS, VALUE_PROP_ENGINEERING,
    VALUE_PROP_BUSINESS, SWOT, REFACTOR_PRIORITIES,
    FINANCIAL, PRICING, SEGMENTS,
)

OUTPUT = "/home/dtfrost/neuralmind/docs/market-research/strategic-report.docx"

# ── helpers ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    # python-docx List Bullet style handles indentation
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.75)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_table(doc, headers, rows, header_color="1A3C6E"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    # header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)
    # data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
    return table


# ── document ─────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── TITLE PAGE ───────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("NeuralMind")
    run.font.size = Pt(44)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Strategic Market & Engineering Report")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("Local-First Graph Intelligence for the AI-Native Enterprise")
    run.font.size = Pt(14)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Prepared: {datetime.now().strftime('%B %Y')}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    src = doc.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = src.add_run("Sources: Business Research Company · Mordor Intelligence · SNS Insider")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    add_page_break(doc)

    # ── TABLE OF CONTENTS ────────────────────────────────────────────────
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Market Analysis",
        "3. Competitive Landscape",
        "4. Future-Proofing Strategy",
        "5. SWOT Analysis",
        "6. Code Quality Findings",
        "7. Refactoring Roadmap",
        "8. Go-to-Market Strategy",
        "9. Financial Projections",
        "10. Conclusion",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(4)

    add_page_break(doc)

    # ── 1. EXECUTIVE SUMMARY ─────────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(doc,
        "NeuralMind is a local-first, graph-based code intelligence platform purpose-built for "
        "engineering teams that cannot — or will not — send proprietary source code to cloud AI services. "
        "In a market projected to grow from $7.65 B in 2025 to $22.2 B by 2030 (CAGR 23.8%), "
        "NeuralMind occupies a defensible niche: privacy-preserving, on-premise AI that indexes code "
        "into a semantic graph and delivers IDE-agnostic intelligence via the Model Context Protocol (MCP).",
        size=11,
    )
    add_para(doc,
        "This report synthesizes market intelligence, competitive analysis, internal code-quality findings, "
        "and a five-year financial model to chart the path from open-source community project to "
        "category-leading enterprise product.",
        size=11,
    )

    add_heading(doc, "Key Takeaways", level=2)
    takeaways = [
        "The AI code tools market is growing at 23.8% CAGR — $7.65 B → $22.2 B by 2030 (Business Research Company, Mordor Intelligence, SNS Insider).",
        "91% of developers already use AI coding tools; 41% of generated code is AI-assisted.",
        "Privacy regulations and enterprise risk posture are driving demand for local-first / air-gapped solutions.",
        "NeuralMind's graph-based semantic intelligence is technically differentiated from LLM-only competitors.",
        "Immediate engineering priority: decompose core.py (1,811 lines) and remove the regex hack in experiment.py.",
        "Revenue model targets $2.4 M ARR by Year 5 with a land-and-expand enterprise motion.",
    ]
    for t in takeaways:
        add_bullet(doc, t)

    add_page_break(doc)

    # ── 2. MARKET ANALYSIS ───────────────────────────────────────────────
    add_heading(doc, "2. Market Analysis", level=1)

    add_heading(doc, "2.1 Total Addressable Market", level=2)
    add_para(doc,
        "The global AI code tools market was valued at $7.65 B in 2025 and is projected to reach "
        f"$22.2 B by 2030, representing a compound annual growth rate of {MARKET['cagr']*100:.1f}%. "
        "This growth is underpinned by near-universal developer adoption and the rapid integration of "
        "AI into every phase of the software development lifecycle.",
        size=11,
    )

    add_table(doc,
        ["Metric", "Value", "Source"],
        [
            ["Market Size 2025", f"${MARKET['current_size_2025']} B", "Business Research Company"],
            ["Market Size 2030", f"${MARKET['projected_2030']} B", "Mordor Intelligence"],
            ["CAGR", f"{MARKET['cagr']*100:.1f}%", "SNS Insider"],
            ["Cloud Share 2025", f"{MARKET['cloud_share_2025']*100:.1f}%", "Mordor Intelligence"],
            ["On-Prem CAGR", f"{MARKET['onprem_cagr']*100:.1f}%", "SNS Insider"],
            ["Enterprise AI 2024", f"${MARKET['enterprise_ai_2024']} B", "Gartner"],
            ["Developer Adoption 2025", f"{MARKET['developer_adoption_2025']*100:.0f}%", "Industry surveys"],
            ["AI-Generated Code", f"{MARKET['code_generation_pct']*100:.0f}%", "GitHub / McKinsey"],
            ["Gartner 2028 Adoption", f"{MARKET['gartner_2028_developer_adoption']*100:.0f}%", "Gartner"],
            ["AI Project Cancellation 2027", f"{MARKET['ai_project_cancellation_2027']*100:.0f}%", "Industry analysis"],
        ],
    )

    add_heading(doc, "2.2 Key Market Drivers", level=2)
    drivers = [
        "Regulatory pressure: GDPR, HIPAA, ITAR, and emerging AI-specific regulations make cloud code analysis legally complex.",
        "Enterprise risk: 40% of AI projects are projected to be cancelled by 2027 due to governance failures — local-first mitigates this.",
        "Developer demand: 91% of developers already use AI tools; the remaining 9% represent a compliance-constrained market.",
        "Graph RAG unbundling: Code intelligence is decoupling from the IDE, creating room for standalone platforms.",
        "Multi-language complexity: Modern codebases span 5+ languages; tree-sitter / SCIP-based parsing is a moat.",
    ]
    for d in drivers:
        add_bullet(doc, d)

    add_heading(doc, "2.3 Market Constraints", level=2)
    constraints = [
        "GitHub Copilot's network effects and brand recognition.",
        "Cloud AI's zero-friction onboarding vs. on-prem deployment complexity.",
        "Open-source alternatives (Continue.dev) with no revenue model.",
        "Rapid commoditization of LLM capabilities reducing differentiation.",
    ]
    for c in constraints:
        add_bullet(doc, c)

    add_page_break(doc)

    # ── 3. COMPETITIVE LANDSCAPE ─────────────────────────────────────────
    add_heading(doc, "3. Competitive Landscape", level=1)
    add_para(doc,
        "The AI code tools market spans cloud IDE assistants, enterprise code intelligence platforms, "
        "and open-source alternatives. NeuralMind competes most directly with the on-prem / privacy-focused "
        "segment while offering a differentiated graph-intelligence core.",
        size=11,
    )

    add_table(doc,
        ["Competitor", "Type", "Strength", "Weakness", "Pricing"],
        [
            [c["name"], c["type"], c["strength"], c["weakness"], c["pricing"]]
            for c in COMPETITORS
        ],
    )

    add_heading(doc, "3.1 Competitive Positioning", level=2)
    add_para(doc,
        "NeuralMind is the only local-first platform that combines graph-based semantic indexing, "
        "self-learning synapses, and pluggable vector backends. Unlike GitHub Copilot (cloud-locked) "
        "or Tabnine (model-only), NeuralMind treats the codebase as a queryable knowledge graph — "
        "enabling multi-hop reasoning, impact analysis, and compliance-grade audit trails.",
        size=11,
    )

    add_page_break(doc)

    # ── 4. FUTURE-PROOFING STRATEGY ──────────────────────────────────────
    add_heading(doc, "4. Future-Proofing Strategy", level=1)
    add_para(doc,
        "To remain relevant as the AI tooling landscape evolves, NeuralMind is architected around "
        "nine strategic pillars that ensure adaptability, compliance, and technical differentiation:",
        size=11,
    )

    for i, pillar in enumerate(PILLARS, 1):
        add_bullet(doc, f"Pillar {i}: {pillar}")

    add_heading(doc, "4.1 Engineering Value Proposition", level=2)
    for v in VALUE_PROP_ENGINEERING:
        add_bullet(doc, v)

    add_heading(doc, "4.2 Business Value Proposition", level=2)
    for v in VALUE_PROP_BUSINESS:
        add_bullet(doc, v)

    add_page_break(doc)

    # ── 5. SWOT ANALYSIS ─────────────────────────────────────────────────
    add_heading(doc, "5. SWOT Analysis", level=1)

    add_table(doc,
        ["", "Factors"],
        [
            ["Strengths", "\n".join(f"• {s}" for s in SWOT["strengths"])],
            ["Weaknesses", "\n".join(f"• {w}" for w in SWOT["weaknesses"])],
            ["Opportunities", "\n".join(f"• {o}" for o in SWOT["opportunities"])],
            ["Threats", "\n".join(f"• {t}" for t in SWOT["threats"])],
        ],
        header_color="2E7D32",
    )

    add_heading(doc, "5.1 Strategic Implications", level=2)
    implications = [
        "Leverage local-first privacy as the primary differentiator against cloud competitors.",
        "Build enterprise sales capability to convert compliance-driven demand.",
        "Invest in IDE integrations to reduce onboarding friction.",
        "Establish thought leadership in graph-based code intelligence.",
    ]
    for imp in implications:
        add_bullet(doc, imp)

    add_page_break(doc)

    # ── 6. CODE QUALITY FINDINGS ─────────────────────────────────────────
    add_heading(doc, "6. Code Quality Findings", level=1)
    add_para(doc,
        "A comprehensive review of the NeuralMind codebase revealed several structural issues that, "
        "if unaddressed, will impede scalability, maintainability, and enterprise readiness. "
        "The findings below are prioritized by business impact.",
        size=11,
    )

    add_heading(doc, "6.1 Critical Findings", level=2)

    add_para(doc, "core.py — 1,811 lines (CRITICAL)", bold=True, size=11)
    add_para(doc,
        "The core module is a monolithic god-object that handles graph construction, query execution, "
        "synapse learning, and server orchestration in a single file. This creates an unmaintainable "
        "change surface where any modification risks cascading regressions. The file must be decomposed "
        "into focused sub-modules: graph_engine, query_planner, synapse_manager, and orchestration.",
        size=11,
    )

    add_para(doc, "ir.py — 963 lines (HIGH)", bold=True, size=11)
    add_para(doc,
        "The intermediate representation module mixes parsing, normalization, and serialization concerns. "
        "While better structured than core.py, it still exceeds the 500-line maintainability threshold. "
        "Recommended split: ir_parser, ir_normalize, ir_serialize.",
        size=11,
    )

    add_para(doc, "experiment.py — Regex Hack (HIGH)", bold=True, size=11)
    add_para(doc,
        "A dangerous regex-based source mutation pattern was identified in experiment.py. This hack "
        "performs unrecoverable in-place modification of source files during experimentation, creating "
        "a data-loss risk and violating the principle of immutable analysis. This code must be removed "
        "immediately and replaced with a sandboxed, AST-based transformation pipeline.",
        size=11,
    )

    add_heading(doc, "6.2 Additional Observations", level=2)
    observations = [
        "server.py (648 lines) mixes HTTP handling, business logic, and telemetry — needs decomposition.",
        "Deprecated enable_reranking attribute is still referenced in 3 call sites — dead code.",
        "Test coverage is below 40% for core modules — refactoring without tests is high-risk.",
        "No CI/CD pipeline enforces linting, type-checking, or security scanning.",
        "Documentation is sparse; API surface is not formally specified.",
    ]
    for obs in observations:
        add_bullet(doc, obs)

    add_page_break(doc)

    # ── 7. REFACTORING ROADMAP ───────────────────────────────────────────
    add_heading(doc, "7. Refactoring Roadmap", level=1)
    add_para(doc,
        "The following prioritized roadmap addresses the code-quality findings above. "
        "Items are ranked by impact-to-effort ratio, with quick wins first.",
        size=11,
    )

    add_table(doc,
        ["Priority", "Item", "Impact", "Effort", "Rationale"],
        [
            [f"P{i+1}", r["item"], r["impact"], r["effort"], r["rationale"]]
            for i, r in enumerate(REFACTOR_PRIORITIES)
        ],
        header_color="C62828",
    )

    add_heading(doc, "7.1 Execution Principles", level=2)
    principles = [
        "Never refactor without tests — write characterization tests first.",
        "Decompose along existing logical boundaries (the code already wants to be split).",
        "Maintain backward compatibility for OSS users during the transition.",
        "Each PR should be independently deployable — no big-bang rewrites.",
    ]
    for pr in principles:
        add_bullet(doc, pr)

    add_page_break(doc)

    # ── 8. GO-TO-MARKET STRATEGY ─────────────────────────────────────────
    add_heading(doc, "8. Go-to-Market Strategy", level=1)
    add_para(doc,
        "NeuralMind's go-to-market motion follows a land-and-expand model targeting four distinct "
        "customer segments. The strategy leverages open-source adoption as a top-of-funnel engine "
        "and converts compliance-driven enterprises into high-value accounts.",
        size=11,
    )

    add_heading(doc, "8.1 Target Segments", level=2)
    add_table(doc,
        ["Segment", "Pain Point", "NeuralMind Fit"],
        [
            [s["name"], s["pain"], s["fit"]]
            for s in SEGMENTS
        ],
    )

    add_heading(doc, "8.2 Pricing Model", level=2)
    add_table(doc,
        ["Tier", "Price", "Features"],
        [
            [k.upper(), v["price"], v["features"]]
            for k, v in PRICING.items()
        ],
    )

    add_heading(doc, "8.3 GTM Motion", level=2)
    gtm_steps = [
        "Phase 1 (Y1): Open-source launch, community building, technical content marketing.",
        "Phase 2 (Y2): First enterprise pilots via compliance-driven inbound (fintech, defense).",
        "Phase 3 (Y3): Team tier launch with self-serve onboarding; hire first enterprise AE.",
        "Phase 4 (Y4): Multi-region expansion, partnership channel (SI/ISV), SOC2 certification.",
        "Phase 5 (Y5): Category leadership in local-first code intelligence; explore strategic options.",
    ]
    for step in gtm_steps:
        add_bullet(doc, step)

    add_page_break(doc)

    # ── 9. FINANCIAL PROJECTIONS ─────────────────────────────────────────
    add_heading(doc, "9. Financial Projections (5-Year)", level=1)
    add_para(doc,
        "The following conservative financial model assumes organic growth driven by open-source adoption "
        "and enterprise conversion. Revenue is net of churn; cost includes engineering, infrastructure, "
        "and initial sales hire.",
        size=11,
    )

    add_table(doc,
        ["Year", "Revenue", "Cost", "Net", "Users", "Notes"],
        [
            [
                f"Year {i+1}",
                f"${y['revenue']:,}",
                f"${y['cost']:,}",
                f"${y['revenue'] - y['cost']:,}",
                f"{y['users']:,}",
                y['notes'],
            ]
            for i, y in enumerate([
                FINANCIAL["year_1"], FINANCIAL["year_2"], FINANCIAL["year_3"],
                FINANCIAL["year_4"], FINANCIAL["year_5"],
            ])
        ],
        header_color="1565C0",
    )

    add_heading(doc, "9.1 Key Assumptions", level=2)
    assumptions = [
        "Year 1 is pure investment — no revenue, community building only.",
        "Enterprise conversion rate: 5% of OSS users convert to paid within 12 months.",
        "Average revenue per user (ARPU): $1,200/year for Team tier, $50K/year for Enterprise.",
        "Headcount grows from 2 (Y1) to 12 (Y5) — engineering-heavy, lean sales.",
        "No external funding assumed — bootstrapped or revenue-financed growth.",
    ]
    for a in assumptions:
        add_bullet(doc, a)

    add_heading(doc, "9.2 Path to Profitability", level=2)
    add_para(doc,
        "The model reaches cash-flow positivity in Year 4 ($330K net) and scales to $1.3 M net "
        "by Year 5. The primary lever is enterprise tier pricing, which carries 80%+ gross margins "
        "once the platform is established.",
        size=11,
    )

    add_page_break(doc)

    # ── 10. CONCLUSION ───────────────────────────────────────────────────
    add_heading(doc, "10. Conclusion", level=1)
    add_para(doc,
        "NeuralMind stands at the intersection of two powerful trends: the explosive growth of AI "
        "code tools ($7.65 B → $22.2B by 2030) and the enterprise demand for privacy-preserving, "
        "local-first AI. The technical differentiation is real — graph-based semantic intelligence "
        "with self-learning capabilities is a genuine moat against LLM-only competitors.",
        size=11,
    )
    add_para(doc,
        "However, the codebase requires immediate attention. The monolithic core.py, the dangerous "
        "regex hack in experiment.py, and the lack of test coverage collectively represent existential "
        "risk to the project's credibility and scalability. The refactoring roadmap outlined in this "
        "report must be executed in parallel with the go-to-market motion — not after.",
        size=11,
    )
    add_para(doc,
        "The path forward is clear: stabilize the codebase, launch the open-source community, convert "
        "compliance-driven enterprises, and build toward category leadership in local-first code "
        "intelligence. With disciplined execution, NeuralMind can capture meaningful share of a "
        "$22 B market while maintaining the privacy-first principles that define its brand.",
        size=11,
    )

    add_heading(doc, "Recommended Immediate Actions", level=2)
    actions = [
        "Remove the regex hack in experiment.py — this week.",
        "Write characterization tests for core.py before any decomposition.",
        "Begin core.py decomposition into graph_engine, query_planner, synapse_manager.",
        "Publish open-source repository with clear README and contribution guidelines.",
        "Initiate SOC2 readiness assessment for Year 4 enterprise readiness.",
        "Hire first enterprise account executive for Year 2 pilot pipeline.",
    ]
    for a in actions:
        add_bullet(doc, a)

    # ── SAVE ─────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f"✓ Strategic report saved to {OUTPUT}")


if __name__ == "__main__":
    build()
